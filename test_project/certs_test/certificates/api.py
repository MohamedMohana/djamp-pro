##############################################################################
# api.py CERTIFICATE GENERATION API
##############################################################################
import hashlib
from rest_framework.decorators import api_view, authentication_classes, permission_classes
from rest_framework.authentication import BasicAuthentication
from rest_framework.permissions import IsAuthenticated
from rest_framework.response import Response
from rest_framework import status
from .forms import CertificateForm
from .certificate_generator import generate_certificate
from .models import Certificate
from django.contrib.auth.models import User
from .email_service import send_email
from datetime import datetime
from .api_auth import APIBasicAuthentication

@api_view(['POST'])
@authentication_classes([BasicAuthentication])
@permission_classes([IsAuthenticated])
def generate_certificate_api(request):
    if request.user.username != 'laravel':
        return Response({'message': 'Unauthorized'}, status=status.HTTP_401_UNAUTHORIZED)

    form = CertificateForm(request.data)
    if form.is_valid():
        template_name = form.cleaned_data.get('template')
        name = form.cleaned_data.get('name')
        course_name = form.cleaned_data.get('course_name')
        duration = form.cleaned_data.get('duration')
        email = form.cleaned_data.get('email')
        send_email_option = form.cleaned_data.get('send_email')
        if isinstance(send_email_option, str):
            send_email_option = send_email_option.lower() == 'true'

        # Include the current timestamp in the hash
        current_timestamp = datetime.now().strftime("%Y%m%d%H%M%S%f")
        hash_input = name + course_name + duration + email + current_timestamp
        hash_value = hashlib.sha256(hash_input.encode('utf-8')).hexdigest()

        output_file = f'media/{hash_value}.pdf'
        output_file_url = request.build_absolute_uri(f'/media/{hash_value}.pdf')

        try:
            generate_certificate(name, course_name, duration, hash_value, template_name, output_file)

            user = User.objects.get(username='laravel')
            Certificate.objects.create(
                name=name,
                email=email,
                course_name=course_name,
                duration=duration,
                hash_value=hash_value,
                certificate_hash=hash_value,
                created_by=user
            )

            if send_email_option:
                subject = "شهادة حضور دورة " + course_name
                email_body = "سعدنا بحضوركم .. ونتطلع لرؤيتكم في دورات قادمة"
                send_email(email, subject, email_body, output_file)

            return Response({'certificate_url': output_file_url}, status=status.HTTP_201_CREATED)
        except Exception as e:
            return Response({'message': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    return Response(form.errors, status=status.HTTP_400_BAD_REQUEST)




##############################################################################
# api.py HASH API
##############################################################################
from rest_framework.decorators import api_view, authentication_classes, permission_classes
from rest_framework.authentication import BasicAuthentication
from rest_framework.permissions import IsAuthenticated
from rest_framework.response import Response
from rest_framework import status
import hashlib
from datetime import datetime
from django.contrib.auth.models import User
from .models import Certificate

ALLOWED_HASH_API_USERS = {'hash_api', 'event_api'}


@api_view(['POST'])
@authentication_classes([APIBasicAuthentication])
@permission_classes([IsAuthenticated])
def generate_hash_api(request):
    if request.user.username not in ALLOWED_HASH_API_USERS:
        return Response({'message': 'Unauthorized'}, status=status.HTTP_401_UNAUTHORIZED)

    # Extract fields
    name = request.data.get('name')
    course_name = request.data.get('course_name')
    email = request.data.get('email')
    duration = request.data.get('duration', '')  # Optional field

    # Validate required fields
    if not all([name, course_name, email]):
        return Response({'message': 'Missing required fields'}, status=status.HTTP_400_BAD_REQUEST)

    # Generate the hash value with timestamp
    current_timestamp = datetime.now().strftime("%Y%m%d%H%M%S%f")
    hash_input = name + course_name + duration + email + current_timestamp
    hash_value = hashlib.sha256(hash_input.encode('utf-8')).hexdigest()

    # Create a new certificate record
    try:
        Certificate.objects.create(
            name=name,
            email=email,
            course_name=course_name,
            duration=duration,
            hash_value=hash_value,
            certificate_hash=hash_value,
            created_by=request.user
        )
    except Exception as e:
        return Response({'message': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    # Create validation URL
    validation_url = f'https://certs.kku.edu.sa/validate/?hash={hash_value}'

    # Return the hash and validation URL
    return Response({'hash': hash_value, 'validation_url': validation_url}, status=status.HTTP_200_OK)


##############################################################################
# api.py KKUX API
##############################################################################

@api_view(['POST'])
@authentication_classes([APIBasicAuthentication])
@permission_classes([IsAuthenticated])
def generate_kkux_api(request):
    # Check if the user is authorized (here, we are checking for the username 'kkux_api')
    if request.user.username != 'kkux_api':
        return Response({'message': 'Unauthorized'}, status=status.HTTP_401_UNAUTHORIZED)

    # Extract fields
    name = request.data.get('name')
    course_name = request.data.get('course_name')
    email = request.data.get('email')
    duration = request.data.get('duration', '')  # Optional field

    # Validate required fields
    if not all([name, course_name, email]):
        return Response({'message': 'Missing required fields'}, status=status.HTTP_400_BAD_REQUEST)

    # Generate the hash value with timestamp
    current_timestamp = datetime.now().strftime("%Y%m%d%H%M%S%f")
    hash_input = name + course_name + duration + email + current_timestamp
    hash_value = hashlib.sha256(hash_input.encode('utf-8')).hexdigest()

    # Create a new certificate record
    try:
        user = User.objects.get(username='kkux_api')  # or use request.user if the logged-in user creates the certificate
        Certificate.objects.create(
            name=name,
            email=email,
            course_name=course_name,
            duration=duration,
            hash_value=hash_value,
            certificate_hash=hash_value,
            created_by=user
        )
    except Exception as e:
        return Response({'message': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    # Create validation URL
    validation_url = f'https://certs.kku.edu.sa/validate/?hash={hash_value}'

    # Return the hash and validation URL
    return Response({'hash': hash_value, 'validation_url': validation_url}, status=status.HTTP_200_OK)





# ##############################################################################
# # api.py BLOCK LISTED API
# # ##############################################################################
# from rest_framework.decorators import api_view, authentication_classes, permission_classes
# from rest_framework.authentication import BasicAuthentication
# from rest_framework.permissions import IsAuthenticated
# from rest_framework.response import Response
# from rest_framework import status
# from django.core.files.storage import default_storage
# from django.conf import settings
# import pandas as pd
# import zipfile
# import xml.etree.ElementTree as ET
# from docx import Document
# import re
# import os

# # Register the namespaces
# namespaces = {
#     'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
# }

# def normalize_arabic(text):
#     if pd.isna(text):
#         return ''
#     text = re.sub(r'[أإآا]', 'ا', text)
#     text = re.sub(r'\sبن\s', ' ', text)
#     text = re.sub(r'[^\w\s]', '', text)
#     text = re.sub(r'ـ', '', text)
#     text = re.sub(r'عبد\s', 'عبد', text)
#     text = re.sub(r'ابو\s', 'ابو', text)
#     text = re.sub(r'ال\s(?!$)', 'ال ', text)  # Add a space if it's not at the end of the text
#     text = re.sub(r'[ء-ي]ّ', '', text)       # Remove shadda
#     text = re.sub(r'\s+', ' ', text)          # Normalize spaces
#     return text.strip()

# def search_in_docx_fields(docx_filename):
#     results = []
#     with zipfile.ZipFile(docx_filename) as docx:
#         for item in docx.filelist:
#             if item.filename not in ['word/document.xml', 'word/footnotes.xml', 'word/endnotes.xml']:
#                 continue
#             with docx.open(item) as file_obj:
#                 xml_content = file_obj.read()
#                 tree = ET.ElementTree(ET.fromstring(xml_content))
#                 root = tree.getroot()
#                 for elem in root.iter():
#                     if elem.tag.endswith('t') or elem.tag.endswith('instrText'):
#                         if elem.text and elem.text.strip():
#                             results.append(elem.text)
#     return results

# def read_docx(file_path):
#     doc = Document(file_path)
#     full_text = []
#     for para in doc.paragraphs:
#         full_text.append(para.text)
#     for table in doc.tables:
#         for row in table.rows:
#             for cell in row.cells:
#                 full_text.append(cell.text)
#     return '\n'.join(full_text)

# def generate_name_variations(name):
#     parts = name.split()
#     if len(parts) == 1:
#         return {
#             'full_name': name
#         }
#     elif len(parts) == 2:
#         return {
#             'full_name': name,
#             'first_and_last': ' '.join([parts[0], parts[1]])
#         }
#     else:
#         return {
#             'full_name': name,
#             'first_and_last': ' '.join([parts[0], parts[-1]]),
#             'last_name': parts[-1],
#             'first_and_second': ' '.join(parts[:2])
#         }

# def strict_match(pattern, text):
#     if not pattern:
#         return False
#     # Escape special characters and match whole words
#     pattern = re.escape(pattern)
#     # Ensure that we're matching full names or significant patterns only
#     return bool(re.search(rf'\b{pattern}\b', text, re.IGNORECASE))

# def check_document_against_blocklist(document_text, xml_texts, block_listed_df):
#     results = []
#     document_text += '\n'.join(xml_texts)  # Combine all document texts
#     document_text_normalized = normalize_arabic(document_text)
    
#     matched_titles = set()  # To avoid duplicating book titles

#     # First, check for **both** book titles and authors, allowing partial name matches if title matches
#     for _, row in block_listed_df.iterrows():
#         book_title = normalize_arabic(row['اسم الكتاب'])
#         author_name = normalize_arabic(row['اسم المؤلف'])
#         author_name_parts = generate_name_variations(author_name)  # Break down the name into variations

#         # Check if the book title is present
#         if strict_match(book_title, document_text_normalized):
#             # Try to match the full author name first
#             if strict_match(author_name, document_text_normalized):
#                 result = {
#                     'author': author_name,
#                     'title': book_title,
#                     'match_confidence': 'full'
#                 }
#                 if result not in results:
#                     results.append(result)
#                     matched_titles.add(book_title)  # Track matched book titles

#             # If full name doesn't match, try partial name matches (e.g., first and last name, or last name only)
#             else:
#                 for variation in author_name_parts.values():
#                     if strict_match(variation, document_text_normalized):
#                         result = {
#                             'author': author_name,
#                             'title': book_title,
#                             'match_confidence': 'partial'  # Mark as a partial match
#                         }
#                         if result not in results:
#                             results.append(result)
#                             matched_titles.add(book_title)  # Track matched book titles

#     return results


# @api_view(['POST'])
# @authentication_classes([BasicAuthentication])
# @permission_classes([IsAuthenticated])
# def block_listed_api(request):
#     # Check if the user is authorized (here, we are checking for the username 'block_listed_api')
#     if request.user.username != 'block_listed_api':
#         return Response({'message': 'Unauthorized'}, status=status.HTTP_401_UNAUTHORIZED)

#     if 'file' not in request.FILES:
#         return Response({'message': 'No file uploaded'}, status=status.HTTP_400_BAD_REQUEST)

#     uploaded_file = request.FILES['file']
#     temp_file_path = 'temp.docx'

#     # Save the uploaded file
#     try:
#         with default_storage.open(temp_file_path, 'wb') as temp_file:
#             for chunk in uploaded_file.chunks():
#                 temp_file.write(chunk)

#         # Verify the file exists
#         if not os.path.exists(default_storage.path(temp_file_path)):
#             return Response({'message': 'File not found after saving'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

#         excel_file_path = settings.BLOCK_LISTED_EXCEL_PATH

#         block_listed_df = pd.read_excel(excel_file_path)

#         document_text = read_docx(default_storage.path(temp_file_path))

#         xml_field_texts = search_in_docx_fields(default_storage.path(temp_file_path))

#         check_results = check_document_against_blocklist(document_text, xml_field_texts, block_listed_df)

#         return Response({
#             'results': check_results,
#             'flag': bool(check_results)
#         }, status=status.HTTP_200_OK)
#     except Exception as e:
#         return Response({'message': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
#     finally:
#         default_storage.delete(temp_file_path)





















































































































#######full and not parial#########
from rest_framework.decorators import api_view, authentication_classes, permission_classes
from rest_framework.authentication import BasicAuthentication
from rest_framework.permissions import IsAuthenticated
from rest_framework.response import Response
from rest_framework import status
from django.core.files.storage import default_storage
from django.conf import settings
import pandas as pd
import zipfile
import xml.etree.ElementTree as ET
from docx import Document
import re
import os

# Register the namespaces
namespaces = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
}

def normalize_arabic(text):
    if pd.isna(text):
        return ''
    text = re.sub(r'[أإآا]', 'ا', text)
    text = re.sub(r'\sبن\s', ' ', text)
    text = re.sub(r'[^\w\s]', '', text)
    text = re.sub(r'ـ', '', text)
    text = re.sub(r'عبد\s', 'عبد', text)
    text = re.sub(r'ابو\s', 'ابو', text)
    text = re.sub(r'ال\s(?!$)', 'ال ', text)  # Add a space if it's not at the end of the text
    text = re.sub(r'[ء-ي]ّ', '', text)       # Remove shadda
    text = re.sub(r'\s+', ' ', text)          # Normalize spaces
    return text.strip()

def search_in_docx_fields(docx_filename):
    results = []
    with zipfile.ZipFile(docx_filename) as docx:
        for item in docx.filelist:
            if item.filename not in ['word/document.xml', 'word/footnotes.xml', 'word/endnotes.xml']:
                continue
            with docx.open(item) as file_obj:
                xml_content = file_obj.read()
                tree = ET.ElementTree(ET.fromstring(xml_content))
                root = tree.getroot()
                for elem in root.iter():
                    if elem.tag.endswith('t') or elem.tag.endswith('instrText'):
                        if elem.text and elem.text.strip():
                            results.append(elem.text)
    return results

def read_docx(file_path):
    doc = Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text.append(cell.text)
    return '\n'.join(full_text)

def generate_name_variations(name):
    parts = name.split()
    if len(parts) == 1:
        return {
            'full_name': name  # Only consider full name if it's a single part
        }
    elif len(parts) == 2:
        return {
            'full_name': name,  # Full name is important
            'first_and_last': ' '.join([parts[0], parts[1]])  # First and last name
        }
    else:
        # Limit the variations to the most meaningful ones
        return {
            'full_name': name,  # Full name is primary
            'first_and_last': ' '.join([parts[0], parts[-1]])  # First and last name
        }

def strict_match(pattern, text):
    if not pattern:
        return False
    # Escape special characters and match whole words
    pattern = re.escape(pattern)
    return bool(re.search(rf'\b{pattern}\b', text, re.IGNORECASE))

def check_proximity(author_name, book_title, document_text):
    author_pos = [m.start() for m in re.finditer(re.escape(author_name), document_text)]
    title_pos = [m.start() for m in re.finditer(re.escape(book_title), document_text)]
    
    # Define proximity threshold (e.g., within 100 characters)
    proximity_threshold = 100
    
    for a_pos in author_pos:
        for t_pos in title_pos:
            if abs(a_pos - t_pos) <= proximity_threshold:
                return True
    return False

def strict_partial_match(pattern, text, min_match_length=3):
    pattern_parts = pattern.split()
    match_count = sum(1 for part in pattern_parts if re.search(rf'\b{re.escape(part)}\b', text, re.IGNORECASE))
    return match_count >= min_match_length

def check_document_against_blocklist(document_text, xml_texts, block_listed_df):
    results = []
    document_text += '\n'.join(xml_texts)  # Combine all document texts
    document_text_normalized = normalize_arabic(document_text)
    
    matched_titles = set()  # To avoid duplicating book titles

    for _, row in block_listed_df.iterrows():
        book_title = normalize_arabic(row['اسم الكتاب'])
        author_name = normalize_arabic(row['اسم المؤلف'])
        author_name_parts = generate_name_variations(author_name)  # Break down the name into variations

        # Check if the book title is present
        if strict_match(book_title, document_text_normalized):
            # Try to match the full author name first
            if strict_match(author_name, document_text_normalized):
                result = {
                    'author': author_name,
                    'title': book_title,
                    'match_confidence': 'full'
                }
                if result not in results:
                    results.append(result)
                    matched_titles.add(book_title)  # Track matched book titles

            # If full name doesn't match, try partial name matches with proximity check
            else:
                for variation in author_name_parts.values():
                    if strict_partial_match(variation, document_text_normalized) and check_proximity(variation, book_title, document_text_normalized):
                        result = {
                            'author': author_name,
                            'title': book_title,
                            'match_confidence': 'partial'  # Mark as a partial match
                        }
                        if result not in results:
                            results.append(result)
                            matched_titles.add(book_title)  # Track matched book titles

    return results


@api_view(['POST'])
@authentication_classes([BasicAuthentication])
@permission_classes([IsAuthenticated])
def block_listed_api(request):
    # Check if the user is authorized (here, we are checking for the username 'block_listed_api')
    if request.user.username != 'block_listed_api':
        return Response({'message': 'Unauthorized'}, status=status.HTTP_401_UNAUTHORIZED)

    if 'file' not in request.FILES:
        return Response({'message': 'No file uploaded'}, status=status.HTTP_400_BAD_REQUEST)

    uploaded_file = request.FILES['file']
    temp_file_path = 'temp.docx'

    # Save the uploaded file
    try:
        with default_storage.open(temp_file_path, 'wb') as temp_file:
            for chunk in uploaded_file.chunks():
                temp_file.write(chunk)

        # Verify the file exists
        if not os.path.exists(default_storage.path(temp_file_path)):
            return Response({'message': 'File not found after saving'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

        excel_file_path = settings.BLOCK_LISTED_EXCEL_PATH

        block_listed_df = pd.read_excel(excel_file_path)

        document_text = read_docx(default_storage.path(temp_file_path))

        xml_field_texts = search_in_docx_fields(default_storage.path(temp_file_path))

        check_results = check_document_against_blocklist(document_text, xml_field_texts, block_listed_df)

        return Response({
            'results': check_results,
            'flag': bool(check_results)
        }, status=status.HTTP_200_OK)
    except Exception as e:
        return Response({'message': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
    finally:
        default_storage.delete(temp_file_path)